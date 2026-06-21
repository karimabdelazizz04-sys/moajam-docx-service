import os
import uuid
import tempfile
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Cm
from copy import deepcopy
from docx import Document
from docx.shared import Cm
from flask import Flask, request, jsonify, send_from_directory


app = Flask(__name__)

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "letterhead_template.docx"
OUTPUT_DIR = BASE_DIR / "generated"
OUTPUT_DIR.mkdir(exist_ok=True)


@app.route("/")
def home():
    return "Moajam LibreOffice HTML to DOCX Service Running"


@app.route("/generate-docx", methods=["POST"])
def generate_docx():
    data = request.get_json(force=True, silent=True) or {}

    job_number = clean_filename(data.get("job_number") or "MOAJAM-JOB")
    final_html = data.get("final_html") or data.get("translated_html") or ""
    translated_text = data.get("translated_text") or ""

    if not final_html and not translated_text:
        return jsonify({"status": "error", "message": "Missing final_html/translated_text"}), 400

    try:
        html = final_html if final_html else text_to_html(translated_text)
        html_path = write_html_file(html)
      converted_docx = convert_html_to_docx(html_path)

filename = f"{job_number}-Final-Translation-{uuid.uuid4().hex[:8]}.docx"
output_path = OUTPUT_DIR / filename

build_final_docx_from_template(converted_docx, output_path)

        base_url = request.host_url.rstrip("/")
        return jsonify({
            "status": "success",
            "download_url": f"{base_url}/download/{filename}",
            "filename": filename
        })

    except Exception as exc:
        return jsonify({"status": "error", "message": str(exc)}), 500


@app.route("/download/<path:filename>")
def download(filename):
    return send_from_directory(str(OUTPUT_DIR), filename, as_attachment=True)


def write_html_file(inner_html):
    full_html = """<!doctype html>
<html lang="ar" dir="rtl">
<head>
<meta charset="utf-8">
<style>
@page {
    size: A4;
    margin-top: 4.5cm;
    margin-bottom: 3.5cm;
    margin-left: 1.8cm;
    margin-right: 1.8cm;
}
html, body {
    direction: rtl;
    text-align: justify;
    font-family: "Sakkal Majalla", "Arial", "Tahoma", sans-serif;
    font-size: 14pt;
    line-height: 1.35;
}
body { margin: 0; }
p {
    direction: rtl;
    text-align: justify;
    margin: 5pt 0;
}
h1 {
    direction: rtl;
    text-align: center;
    font-size: 18pt;
    font-weight: bold;
    margin: 10pt 0 8pt 0;
}
h2 {
    direction: rtl;
    text-align: right;
    font-size: 16pt;
    font-weight: bold;
    color: #1f4e79;
    margin: 10pt 0 6pt 0;
}
h3, h4 {
    direction: rtl;
    text-align: right;
    font-size: 14.5pt;
    font-weight: bold;
    color: #1f4e79;
    margin: 8pt 0 5pt 0;
}
table {
    width: 100%;
    border-collapse: collapse;
    direction: rtl;
    margin: 8pt 0 12pt 0;
}
th, td {
    border: 1px solid #555555;
    padding: 4pt 6pt;
    vertical-align: top;
    text-align: right;
    direction: rtl;
    font-size: 11.5pt;
}
th {
    background-color: #d9eaf7;
    font-weight: bold;
}
ul, ol {
    direction: rtl;
    text-align: right;
}
li { margin-bottom: 3pt; }
.page-break { page-break-before: always; }
</style>
</head>
<body>
""" + str(inner_html) + """
</body>
</html>
"""

    fd, path = tempfile.mkstemp(suffix=".html", dir=str(BASE_DIR))
    with os.fdopen(fd, "w", encoding="utf-8") as f:
        f.write(full_html)
    return Path(path)


def convert_html_to_docx(html_path):
    out_dir = Path(tempfile.mkdtemp(dir=str(BASE_DIR)))

    cmd = [
        "soffice",
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--invisible",
        "--convert-to",
        "docx:MS Word 2007 XML",
        "--outdir",
        str(out_dir),
        str(html_path)
    ]

    result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)

    if result.returncode != 0:
        raise RuntimeError(
            "LibreOffice failed. STDOUT: "
            + (result.stdout or "")
            + " STDERR: "
            + (result.stderr or "")
        )

    candidates = list(out_dir.glob("*"))

    docx_candidates = list(out_dir.glob("*.docx"))
    if docx_candidates:
        return docx_candidates[0]

    raise RuntimeError(
        "LibreOffice did not create DOCX. Files created: "
        + ", ".join([p.name for p in candidates])
        + " STDOUT: "
        + (result.stdout or "")
        + " STDERR: "
        + (result.stderr or "")
    )
def build_final_docx_from_template(content_docx_path, output_path):
    if not TEMPLATE_PATH.exists():
        raise RuntimeError(f"Template not found: {TEMPLATE_PATH}")

    template_doc = Document(str(TEMPLATE_PATH))
    content_doc = Document(str(content_docx_path))

    apply_docx_layout_to_document(template_doc)

    clear_body_keep_sections(template_doc)

    for element in content_doc.element.body:
        if element.tag.endswith('sectPr'):
            continue
        template_doc.element.body.append(deepcopy(element))

    template_doc.save(str(output_path))


def apply_docx_layout_to_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(4.5)
        section.bottom_margin = Cm(3.5)
        section.right_margin = Cm(1.8)
        section.left_margin = Cm(1.8)


def clear_body_keep_sections(doc):
    body = doc.element.body
    sectPr = body.sectPr

    for child in list(body):
        if child is not sectPr:
            body.remove(child)
def apply_docx_layout(docx_path):
    doc = Document(str(docx_path))

    for section in doc.sections:
        section.top_margin = Cm(4.5)
        section.bottom_margin = Cm(3.5)
        section.right_margin = Cm(1.8)
        section.left_margin = Cm(1.8)

    doc.save(str(docx_path))

def text_to_html(text):
    lines = []
    for line in str(text or "").splitlines():
        line = line.strip()
        if line:
            lines.append("<p>" + escape_html(line) + "</p>")
    return '<div dir="rtl" style="direction:rtl;text-align:justify">' + "\n".join(lines) + "</div>"


def escape_html(value):
    return str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def clean_filename(value):
    value = str(value or "MOAJAM").strip()
    safe = "".join(ch for ch in value if ch.isalnum() or ch in ("-", "_"))
    return safe or "MOAJAM"


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
